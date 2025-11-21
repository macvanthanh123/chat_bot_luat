import json
import re
from docx import Document
from datetime import datetime
from app.utils.logger import logger  

class DocParser:
    def __init__(self, filepath):
        self.filepath = filepath
        logger.info("Bắt đầu phân tích tài liệu: {}", filepath)

        self.paragraphs = self._read_docx_text()
        self.title = self._extract_title()
        self.date = self._extract_date()
        self.text = "\n".join(self.paragraphs)
        self.markdown = self._convert_to_markdown_structured()
        self.url = filepath
        self.images = []

        logger.info("Đã phân tích xong tài liệu | title='{}' | tổng đoạn={}", self.title, len(self.paragraphs))

    def _read_docx_text(self):
        try:
            doc = Document(self.filepath)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            logger.debug("Đã đọc {} đoạn văn bản từ file DOCX", len(paragraphs))
            return paragraphs
        except Exception as e:
            logger.exception("Lỗi khi đọc file DOCX: {}", e)
            return []

    def _extract_title(self):
        for p in self.paragraphs:
            if len(p) > 10:
                return p
        logger.warning("Không tìm thấy tiêu đề rõ ràng trong tài liệu")
        return "Không rõ tiêu đề"

    def _extract_date(self):
        now = datetime.now()
        date_str = f"ngày {now.day:02d} tháng {now.month:02d} năm {now.year}"
        logger.debug("Gán ngày hiện tại cho tài liệu: {}", date_str)
        return date_str

    def _convert_to_markdown_structured(self):
        md_lines = []
        for para in self.paragraphs:
            if "http" in para and (".jpg" in para or ".png" in para):
                md_lines.append(f"![image]({para})")
            elif para.lower().startswith("chương "):
                md_lines.append(f"# {para}")
            elif para.lower().startswith("điều "):
                md_lines.append(f"## {para}")
            elif re.match(r"^(Khoản|khoản|Điểm|điểm|Mục|mục)\s+\d+[\.:)]", para):
                md_lines.append(f"- {para}")
            else:
                md_lines.append(para)

        logger.debug("Đã chuyển đổi markdown với {} dòng", len(md_lines))
        return "\n\n".join(md_lines)

    def to_dict(self):
        return {
            "url": self.url,
            "title": self.title,
            "date": self.date,
            "text": self.text,
            "markdown": self.markdown,
            "images": self.images,
        }

    def save_json(self, filename="vanban.json"):
        try:
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(self.to_dict(), f, ensure_ascii=False, indent=4)
            logger.info("Đã lưu dữ liệu tài liệu thành JSON tại '{}'", filename)
        except Exception as e:
            logger.exception("Lỗi khi lưu JSON: {}", e)

    def save_markdown(self, filename="temp.md"):
        try:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(self.markdown)
            logger.info("Đã lưu markdown ra file '{}'", filename)
        except Exception as e:
            logger.exception("Lỗi khi lưu markdown: {}", e)
