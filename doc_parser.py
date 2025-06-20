import json
import re
from docx import Document


class DocParser:
    def __init__(self, filepath):
        self.filepath = filepath
        self.paragraphs = self._read_docx_text()
        self.title = self._extract_title()
        self.date = self._extract_date()
        self.text = "\n".join(self.paragraphs)
        self.markdown = self._convert_to_markdown_structured()
        self.url = filepath  
        self.images = []  

    def _read_docx_text(self):
        doc = Document(self.filepath)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    def _extract_title(self):
        for p in self.paragraphs:
            if len(p) > 10:
                return p
        return "Không rõ tiêu đề"

    def _extract_date(self):
        pattern = r"ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}"
        for p in self.paragraphs:
            match = re.search(pattern, p, re.IGNORECASE)
            if match:
                return match.group()
        return ""

    def _convert_to_markdown_structured(self):
        md_lines = []
        for para in self.paragraphs:
            if "http" in para and (".jpg" in para or ".png" in para):
                md_lines.append(f"![image]({para})")
            elif para.lower().startswith("chương "):
                md_lines.append(f"# {para}")
            elif para.lower().startswith("điều "):
                md_lines.append(f"## {para}")
            elif re.match(r"^(Khoản|khoản|Điểm|điểm|Mục|mục)\s+\d+[\.:)]", para):
                md_lines.append(f"- {para}")
            else:
                md_lines.append(para)
        return "\n\n".join(md_lines)

    def to_dict(self):
        return {
            "url": self.url,
            "title": self.title,
            "date": self.date,
            "text": self.text,
            "markdown": self.markdown,
            "images": self.images,
        }

    def save_json(self, filename="vanban.json"):
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(self.to_dict(), f, ensure_ascii=False, indent=4)

    def save_markdown(self, filename="temp.md"):
        with open(filename, "w", encoding="utf-8") as f:
            f.write(self.markdown)
